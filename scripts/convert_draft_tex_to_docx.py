import csv
import re
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Inches, Pt, RGBColor
from PIL import Image, ImageDraw, ImageFont


ROOT = Path(__file__).resolve().parents[1]
TEX_PATH = ROOT / "draft.tex"
OUTPUT_PATH = ROOT / "draft.docx"
FIGURES_DIR = ROOT / "results" / "article_figures"


COLORS = {
    "original": (34, 34, 34),
    "jpeg": (217, 95, 2),
    "bpc": (27, 158, 119),
    "map50": (117, 112, 179),
    "map50_95": (102, 166, 30),
    "grid": (220, 220, 220),
    "axis": (45, 45, 45),
}


def load_font(size, bold=False):
    candidates = [
        "arialbd.ttf" if bold else "arial.ttf",
        "calibrib.ttf" if bold else "calibri.ttf",
    ]
    for candidate in candidates:
        try:
            return ImageFont.truetype(candidate, size)
        except OSError:
            continue
    return ImageFont.load_default()


def read_rows():
    metrics_path = FIGURES_DIR / "article_metrics_table.csv"
    with metrics_path.open("r", encoding="utf-8") as file:
        rows = list(csv.DictReader(file))
    for row in rows:
        for key in ["size_mb", "compression_ratio", "map50", "map50_95", "map50_drop_pct", "psnr", "ssim"]:
            row[key] = float(row[key]) if row[key] else None
    return rows


def by_type(rows, kind):
    return [row for row in rows if row["type"] == kind]


def nice_range(values):
    minimum = min(values)
    maximum = max(values)
    if minimum == maximum:
        return minimum - 1, maximum + 1
    span = maximum - minimum
    return minimum - span * 0.05, maximum + span * 0.08


def scale(value, minimum, maximum, start, end, invert=False):
    position = (value - minimum) / (maximum - minimum)
    if invert:
        position = 1 - position
    return start + position * (end - start)


def draw_axes(draw, title, xlabel, ylabel, x_min, x_max, y_min, y_max, width, height):
    title_font = load_font(34, bold=True)
    label_font = load_font(22)
    tick_font = load_font(18)
    left, top, right, bottom = 140, 110, width - 70, height - 140

    draw.rectangle((0, 0, width, height), fill="white")
    draw.text((width / 2, 35), title, font=title_font, fill=(20, 20, 20), anchor="ma")

    for index in range(5):
        x_value = x_min + (x_max - x_min) * index / 4
        x = scale(x_value, x_min, x_max, left, right)
        draw.line((x, top, x, bottom), fill=COLORS["grid"], width=1)
        draw.text((x, bottom + 22), f"{x_value:.2f}", font=tick_font, fill=(40, 40, 40), anchor="ma")

        y_value = y_min + (y_max - y_min) * index / 4
        y = scale(y_value, y_min, y_max, top, bottom, invert=True)
        draw.line((left, y, right, y), fill=COLORS["grid"], width=1)
        draw.text((left - 18, y), f"{y_value:.3f}", font=tick_font, fill=(40, 40, 40), anchor="rm")

    draw.line((left, bottom, right, bottom), fill=COLORS["axis"], width=3)
    draw.line((left, top, left, bottom), fill=COLORS["axis"], width=3)
    draw.text(((left + right) / 2, height - 42), xlabel, font=label_font, fill=(20, 20, 20), anchor="ma")

    label_image = Image.new("RGBA", (420, 45), (255, 255, 255, 0))
    label_draw = ImageDraw.Draw(label_image)
    label_draw.text((210, 22), ylabel, font=label_font, fill=(20, 20, 20), anchor="mm")
    label_image = label_image.rotate(90, expand=True)
    return left, top, right, bottom, label_image


def save_line_chart(rows, path, title, xlabel, ylabel, x_key, y_key, groups):
    width, height = 1800, 1200
    image = Image.new("RGB", (width, height), "white")
    draw = ImageDraw.Draw(image)
    selected = [row for group in groups.values() for row in group]
    x_min, x_max = nice_range([row[x_key] for row in selected])
    y_min, y_max = nice_range([row[y_key] for row in selected])
    left, top, right, bottom, ylabel_image = draw_axes(
        draw, title, xlabel, ylabel, x_min, x_max, y_min, y_max, width, height
    )
    image.paste(ylabel_image, (35, int((top + bottom) / 2 - ylabel_image.height / 2)), ylabel_image)
    text_font = load_font(20)
    legend_font = load_font(20)

    for legend_index, (kind, group) in enumerate(groups.items()):
        color = COLORS[kind]
        points = [
            (
                scale(row[x_key], x_min, x_max, left, right),
                scale(row[y_key], y_min, y_max, top, bottom, invert=True),
            )
            for row in group
        ]
        if len(points) > 1:
            draw.line(points, fill=color, width=5)
        for (x, y), row in zip(points, group):
            draw.ellipse((x - 9, y - 9, x + 9, y + 9), fill=color)
            draw.text((x + 13, y - 18), row["variant"], font=text_font, fill=(20, 20, 20))
        legend_x = left + 30 + legend_index * 160
        legend_y = top - 45
        draw.rectangle((legend_x, legend_y, legend_x + 24, legend_y + 24), fill=color)
        draw.text((legend_x + 34, legend_y + 2), kind.upper(), font=legend_font, fill=(20, 20, 20))

    image.save(path)


def save_scatter(rows, path, title, xlabel, ylabel, x_key, y_key):
    groups = {"jpeg": by_type(rows, "jpeg"), "bpc": by_type(rows, "bpc")}
    selected = [row for group in groups.values() for row in group if row[x_key] is not None]
    save_line_chart(selected, path, title, xlabel, ylabel, x_key, y_key, groups)


def save_size_bars(rows, path):
    ordered = sorted(rows, key=lambda row: row["size_mb"])
    save_simple_bars(ordered, path, "Dataset size by variant", "Dataset size (MB)", "size_mb")


def save_simple_bars(rows, path, title, ylabel, value_key):
    width, height = 1900, 1200
    image = Image.new("RGB", (width, height), "white")
    draw = ImageDraw.Draw(image)
    title_font = load_font(34, bold=True)
    tick_font = load_font(19)
    label_font = load_font(22)
    left, top, right, bottom = 140, 110, width - 70, height - 150
    max_value = max(row[value_key] for row in rows) * 1.12
    draw.text((width / 2, 35), title, font=title_font, fill=(20, 20, 20), anchor="ma")
    for index in range(5):
        value = max_value * index / 4
        y = bottom - (value / max_value) * (bottom - top)
        draw.line((left, y, right, y), fill=COLORS["grid"], width=1)
        draw.text((left - 18, y), f"{value:.2f}", font=tick_font, fill=(40, 40, 40), anchor="rm")
    gap = 18
    bar_width = (right - left - gap * (len(rows) - 1)) / len(rows)
    for index, row in enumerate(rows):
        x = left + index * (bar_width + gap)
        height_value = row[value_key] / max_value * (bottom - top)
        y = bottom - height_value
        draw.rectangle((x, y, x + bar_width, bottom), fill=COLORS[row["type"]])
        draw.text((x + bar_width / 2, bottom + 28), row["variant"], font=tick_font, fill=(20, 20, 20), anchor="ma")
    draw.line((left, bottom, right, bottom), fill=COLORS["axis"], width=3)
    draw.line((left, top, left, bottom), fill=COLORS["axis"], width=3)
    draw.text((width / 2, height - 45), "Variant", font=label_font, fill=(20, 20, 20), anchor="ma")
    image.save(path)


def save_detection_bars(rows, path):
    ordered_names = ["original", "q94", "q88", "q75", "q50", "q25", "b7", "b4", "b3", "b2", "b1"]
    row_map = {row["variant"]: row for row in rows}
    ordered = [row_map[name] for name in ordered_names if name in row_map]
    width, height = 1900, 1200
    image = Image.new("RGB", (width, height), "white")
    draw = ImageDraw.Draw(image)
    title_font = load_font(34, bold=True)
    tick_font = load_font(19)
    label_font = load_font(22)
    left, top, right, bottom = 140, 110, width - 70, height - 150
    max_value = max(max(row["map50"], row["map50_95"]) for row in ordered) * 1.12
    draw.text((width / 2, 35), "Detection metrics by variant", font=title_font, fill=(20, 20, 20), anchor="ma")
    for index in range(5):
        value = max_value * index / 4
        y = bottom - (value / max_value) * (bottom - top)
        draw.line((left, y, right, y), fill=COLORS["grid"], width=1)
        draw.text((left - 18, y), f"{value:.2f}", font=tick_font, fill=(40, 40, 40), anchor="rm")
    group_width = (right - left) / len(ordered)
    bar_width = group_width * 0.32
    for index, row in enumerate(ordered):
        center = left + group_width * index + group_width / 2
        for offset, key, color in [(-bar_width / 2, "map50", COLORS["map50"]), (bar_width / 2, "map50_95", COLORS["map50_95"])]:
            value_height = row[key] / max_value * (bottom - top)
            x = center + offset - bar_width / 2
            y = bottom - value_height
            draw.rectangle((x, y, x + bar_width, bottom), fill=color)
        draw.text((center, bottom + 28), row["variant"], font=tick_font, fill=(20, 20, 20), anchor="ma")
    draw.rectangle((left + 20, top - 45, left + 44, top - 21), fill=COLORS["map50"])
    draw.text((left + 55, top - 42), "mAP50", font=tick_font, fill=(20, 20, 20))
    draw.rectangle((left + 160, top - 45, left + 184, top - 21), fill=COLORS["map50_95"])
    draw.text((left + 195, top - 42), "mAP50-95", font=tick_font, fill=(20, 20, 20))
    draw.line((left, bottom, right, bottom), fill=COLORS["axis"], width=3)
    draw.line((left, top, left, bottom), fill=COLORS["axis"], width=3)
    draw.text((width / 2, height - 45), "Variant", font=label_font, fill=(20, 20, 20), anchor="ma")
    image.save(path)


def ensure_png_figures():
    rows = read_rows()
    pngs = {
        "fig1_map50_vs_size.png": lambda: save_line_chart(
            rows,
            FIGURES_DIR / "fig1_map50_vs_size.png",
            "Detection quality vs dataset size",
            "Dataset size (MB)",
            "mAP50",
            "size_mb",
            "map50",
            {
                "original": by_type(rows, "original"),
                "jpeg": sorted(by_type(rows, "jpeg"), key=lambda row: row["size_mb"]),
                "bpc": sorted(by_type(rows, "bpc"), key=lambda row: row["size_mb"]),
            },
        ),
        "fig2_compression_ratio_vs_map50.png": lambda: save_line_chart(
            [row for row in rows if row["type"] != "original"],
            FIGURES_DIR / "fig2_compression_ratio_vs_map50.png",
            "Compression ratio vs detection quality",
            "Compression ratio vs original",
            "mAP50",
            "compression_ratio",
            "map50",
            {
                "jpeg": sorted(by_type(rows, "jpeg"), key=lambda row: row["compression_ratio"]),
                "bpc": sorted(by_type(rows, "bpc"), key=lambda row: row["compression_ratio"]),
            },
        ),
        "fig3_relative_map50_drop.png": lambda: save_line_chart(
            [row for row in rows if row["type"] != "original"],
            FIGURES_DIR / "fig3_relative_map50_drop.png",
            "Detection loss under compression",
            "Compression ratio vs original",
            "Relative mAP50 drop (%)",
            "compression_ratio",
            "map50_drop_pct",
            {
                "jpeg": sorted(by_type(rows, "jpeg"), key=lambda row: row["compression_ratio"]),
                "bpc": sorted(by_type(rows, "bpc"), key=lambda row: row["compression_ratio"]),
            },
        ),
        "fig4_dataset_size_bars.png": lambda: save_size_bars(rows, FIGURES_DIR / "fig4_dataset_size_bars.png"),
        "fig5_detection_metric_bars.png": lambda: save_detection_bars(rows, FIGURES_DIR / "fig5_detection_metric_bars.png"),
        "fig6_psnr_vs_map50.png": lambda: save_scatter(
            [row for row in rows if row["type"] != "original"],
            FIGURES_DIR / "fig6_psnr_vs_map50.png",
            "PSNR vs detection quality",
            "Mean PSNR",
            "mAP50",
            "psnr",
            "map50",
        ),
        "fig7_ssim_vs_map50.png": lambda: save_scatter(
            [row for row in rows if row["type"] != "original"],
            FIGURES_DIR / "fig7_ssim_vs_map50.png",
            "SSIM vs detection quality",
            "Mean SSIM",
            "mAP50",
            "ssim",
            "map50",
        ),
    }
    for name, builder in pngs.items():
        if not (FIGURES_DIR / name).exists():
            builder()


def extract_env(text, name):
    match = re.search(rf"\\begin\{{{name}\}}(.*?)\\end\{{{name}\}}", text, flags=re.S)
    return match.group(1).strip() if match else ""


def latex_to_text(value):
    replacements = {
        r"\&": "&",
        r"\%": "%",
        r"\_": "_",
        r"\times": "×",
        r"\le": "≤",
        r"\ge": "≥",
        r"\Delta": "Δ",
        r"\textasciitilde": "~",
        r"``": '"',
        r"''": '"',
        "~": " ",
    }
    for old, new in replacements.items():
        value = value.replace(old, new)
    value = re.sub(r"\\cite\{([^}]+)\}", r"[\1]", value)
    value = re.sub(r"\\ref\{([^}]+)\}", r"\1", value)
    value = re.sub(r"\\label\{[^}]+\}", "", value)
    value = re.sub(r"\\textbf\{([^{}]*)\}", r"\1", value)
    value = re.sub(r"\\textit\{([^{}]*)\}", r"\1", value)
    value = re.sub(r"\\texttt\{([^{}]*)\}", r"\1", value)
    value = re.sub(r"\\text\{([^{}]*)\}", r"\1", value)
    value = re.sub(r"\\[a-zA-Z]+\*?(?:\[[^\]]*\])?(?:\{([^{}]*)\})?", lambda m: m.group(1) or "", value)
    value = value.replace("{", "").replace("}", "")
    value = re.sub(r"\s+", " ", value).strip()
    return value


def split_top_level(text):
    pattern = re.compile(
        r"(\\section\{.*?\}|\\subsection\{.*?\}|\\subsubsection\{.*?\}|"
        r"\\begin\{table\*?\}.*?\\end\{table\*?\}|"
        r"\\begin\{figure\*?\}.*?\\end\{figure\*?\}|"
        r"\\begin\{itemize\}.*?\\end\{itemize\}|"
        r"\\begin\{enumerate\}.*?\\end\{enumerate\}|"
        r"\\begin\{equation\}.*?\\end\{equation\}|"
        r"\\begin\{thebibliography\}.*?\\end\{thebibliography\})",
        flags=re.S,
    )
    position = 0
    for match in pattern.finditer(text):
        if match.start() > position:
            yield ("text", text[position : match.start()])
        token = match.group(0)
        if token.startswith(r"\section"):
            yield ("section", token)
        elif token.startswith(r"\subsection"):
            yield ("subsection", token)
        elif token.startswith(r"\subsubsection"):
            yield ("subsubsection", token)
        elif token.startswith(r"\begin{table"):
            yield ("table", token)
        elif token.startswith(r"\begin{figure"):
            yield ("figure", token)
        elif token.startswith(r"\begin{itemize"):
            yield ("itemize", token)
        elif token.startswith(r"\begin{enumerate"):
            yield ("enumerate", token)
        elif token.startswith(r"\begin{equation"):
            yield ("equation", token)
        elif token.startswith(r"\begin{thebibliography"):
            yield ("bibliography", token)
        position = match.end()
    if position < len(text):
        yield ("text", text[position:])


def add_paragraphs(document, text):
    text = re.sub(r"%.*", "", text)
    chunks = [chunk.strip() for chunk in re.split(r"\n\s*\n", text) if chunk.strip()]
    for chunk in chunks:
        if chunk.startswith("\\") and len(chunk.split()) <= 2:
            continue
        paragraph = latex_to_text(chunk)
        if paragraph:
            document.add_paragraph(paragraph)


def add_list(document, token, style):
    items = re.findall(r"\\item\s+(.*?)(?=\\item|\\end\{(?:itemize|enumerate)\})", token, flags=re.S)
    for item in items:
        text = latex_to_text(item)
        if text:
            document.add_paragraph(text, style=style)


def resolve_graphic_path(raw_path):
    path = Path(raw_path)
    candidates = []
    if path.is_absolute():
        candidates.append(path)
    else:
        candidates.append(ROOT / path)
    for candidate in list(candidates):
        stem = candidate.with_suffix("")
        candidates.extend([stem.with_suffix(".png"), stem.with_suffix(".jpg"), stem.with_suffix(".jpeg")])
    for candidate in candidates:
        if candidate.exists():
            return candidate
    return None


def add_figure(document, token):
    graphic_match = re.search(r"\\includegraphics(?:\[[^\]]*\])?\{([^}]+)\}", token)
    caption_match = re.search(r"\\caption\{(.*?)\}", token, flags=re.S)
    image_path = resolve_graphic_path(graphic_match.group(1)) if graphic_match else None
    if image_path:
        paragraph = document.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = paragraph.add_run()
        run.add_picture(str(image_path), width=Inches(6.2))
    if caption_match:
        caption = document.add_paragraph(latex_to_text("Fig. " + caption_match.group(1)))
        caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
        caption.runs[0].italic = True


def add_equation(document, token):
    equation = re.sub(r"\\begin\{equation\}|\\end\{equation\}", "", token).strip()
    equation = re.sub(r"\\label\{[^}]+\}", "", equation).strip()
    paragraph = document.add_paragraph(latex_to_text(equation))
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in paragraph.runs:
        run.italic = True


def add_results_table(document):
    rows = read_rows()
    headers = ["Variant", "Type", "Size MB", "CR", "mAP50", "mAP50-95", "PSNR", "SSIM", "Δ mAP50"]
    table = document.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    header_cells = table.rows[0].cells
    for index, header in enumerate(headers):
        header_cells[index].text = header
    for row in rows:
        cells = table.add_row().cells
        values = [
            row["variant"],
            row["type"],
            f"{row['size_mb']:.4f}",
            f"{row['compression_ratio']:.2f}x",
            f"{row['map50']:.4f}",
            f"{row['map50_95']:.4f}",
            "—" if row["psnr"] is None else f"{row['psnr']:.4f}",
            "—" if row["ssim"] is None else f"{row['ssim']:.4f}",
            f"{row['map50_drop_pct']:.2f}%",
        ]
        for index, value in enumerate(values):
            cells[index].text = value
    for table_row in table.rows:
        for cell in table_row.cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(7)
    caption = document.add_paragraph(
        "Table 1. Comprehensive quantitative results for original, JPEG, and BPC variants on COCO val2017 using YOLOv8n."
    )
    caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption.runs[0].italic = True


def add_table(document, token):
    caption_match = re.search(r"\\caption\{(.*?)\}", token, flags=re.S)
    add_results_table(document)
    if caption_match:
        return


def add_bibliography(document, token):
    document.add_heading("References", level=1)
    items = re.findall(r"\\bibitem\{([^}]+)\}(.*?)(?=\\bibitem|\\end\{thebibliography\})", token, flags=re.S)
    for index, (_, body) in enumerate(items, start=1):
        paragraph = document.add_paragraph(style="List Number")
        paragraph.add_run(latex_to_text(body))


def set_document_style(document):
    section = document.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(3)
    section.right_margin = Cm(1.5)
    style = document.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(12)
    style.paragraph_format.line_spacing = 1
    style.paragraph_format.first_line_indent = Cm(1.25)
    for style_name in ["Heading 1", "Heading 2", "Heading 3"]:
        heading = document.styles[style_name]
        heading.font.name = "Times New Roman"
        heading.font.color.rgb = RGBColor(0, 0, 0)


def build_docx():
    ensure_png_figures()
    tex = TEX_PATH.read_text(encoding="utf-8")
    tex = tex.replace(
        r"\subsection{Analysis of JPEG Sweet Spots" + "\n",
        r"\subsection{Analysis of JPEG Sweet Spots}" + "\n",
    )
    document = Document()
    set_document_style(document)

    title = re.search(r"\\title\{(.*?)\}", tex, flags=re.S)
    author = re.search(r"\\author\{(.*?)\}", tex, flags=re.S)
    if title:
        paragraph = document.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = paragraph.add_run(latex_to_text(title.group(1)))
        run.bold = True
        run.font.size = Pt(14)
    if author:
        paragraph = document.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph.add_run(latex_to_text(author.group(1)))

    abstract = extract_env(tex, "abstract")
    if abstract:
        document.add_heading("Abstract", level=1)
        add_paragraphs(document, abstract)
    keyword_match = re.search(r"\\noindent\s+\\textbf\{Keywords:\}(.*?)(?=\\section)", tex, flags=re.S)
    if keyword_match:
        paragraph = document.add_paragraph()
        run = paragraph.add_run("Keywords: ")
        run.bold = True
        paragraph.add_run(latex_to_text(keyword_match.group(1)))

    document_body = re.search(r"\\begin\{document\}(.*?)\\end\{document\}", tex, flags=re.S)
    body_source = document_body.group(1) if document_body else tex
    body_match = re.search(r"\\section\{", body_source)
    body = body_source[body_match.start() :] if body_match else body_source
    for kind, token in split_top_level(body):
        if kind == "text":
            add_paragraphs(document, token)
        elif kind == "section":
            document.add_heading(latex_to_text(re.search(r"\{(.*?)\}", token, flags=re.S).group(1)), level=1)
        elif kind == "subsection":
            document.add_heading(latex_to_text(re.search(r"\{(.*?)\}", token, flags=re.S).group(1)), level=2)
        elif kind == "subsubsection":
            document.add_heading(latex_to_text(re.search(r"\{(.*?)\}", token, flags=re.S).group(1)), level=3)
        elif kind == "itemize":
            add_list(document, token, "List Bullet")
        elif kind == "enumerate":
            add_list(document, token, "List Number")
        elif kind == "equation":
            add_equation(document, token)
        elif kind == "table":
            add_table(document, token)
        elif kind == "figure":
            add_figure(document, token)
        elif kind == "bibliography":
            add_bibliography(document, token)

    document.save(OUTPUT_PATH)
    print(f"Saved {OUTPUT_PATH}")


if __name__ == "__main__":
    build_docx()
